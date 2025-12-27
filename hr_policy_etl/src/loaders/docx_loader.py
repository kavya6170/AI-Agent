import os
import docx

def load_docx(file_path: str) -> str:
    """
    Load raw text from a .docx file.

    Args:
        file_path (str): The absolute path to the .docx file.

    Returns:
        str: The raw extracted text, with paragraphs separated by newlines.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a .docx file.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.docx'):
        raise ValueError(f"Invalid file extension for DOCX loader: {file_path}")

    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
        
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_text))

    return '\n'.join(full_text)
